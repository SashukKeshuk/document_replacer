from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import locale

def fill_contract_mr(uid, a):
    locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")
    current_date = datetime.datetime.now()
    formatted_date = current_date.strftime('«%d» %B %Y г.')
    text_list = []
    doc = Document("templates/mutual_ur.docx")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith('1.1'):
            p_index = i
            break
    for i, text in enumerate(text_list, start=1):
        new_paragraph = doc.add_paragraph()
        new_paragraph.text = f'1.1.{i} {text}'
        p = new_paragraph._element
        doc.element.body.insert(p_index+i+1, p)
    data = {
    "№______": a[0],
    "«___»________202_ г.": formatted_date,
    "ТОО «и_____________»": f"ТОО «{a[1]}»",
    "БИН:и____________": f"БИН: {a[2]}",
    "директора______________________": f"директора {a[3]}",
    "ТОО «._____________»": f"ТОО «{a[27]}»",
    "БИН:._____________": f"БИН: {a[28]}",
    "директора.______________________": f"директора {a[25]}",
    "основании ____________________": f"основании {a[4]}",
    "адрес: ______________.": f"адрес: {a[9]}",
    "e-mail: _____________.": f"e-mail: {a[10]}",
    "Тел./факс: ____________.": f"Тел./факс: {a[11]}",
    "IBAN: ___________ .": f"IBAN: {a[12]}",
    "в АО __________ .": f"в АО {a[13]}",
    "БИК _____________.": f"БИК {a[14]}",
    "адрес: ______________": f"адрес: {a[15]}",
    "e-mail: _____________": f"e-mail: {a[16]}",
    "Тел./ WhatsApp: ____________": f"Тел./ WhatsApp: {a[17]}",
    "БИН ____________": f"БИН {a[28]}",
    "IBAN: ___________": f"IBAN {a[18]}",
    "в АО __________": f"AO {a[19]}",
    "БИК _____________": f"БИК {a[20]}",
    ", ______________________________.": f", {a[21]}",
    "в течение ____ (_________________)": f"в течение {a[22]}",
    "director_dealer": a[24],
    "director_buyer": a[26],
    }

    for paragraph in doc.paragraphs:
        for key in data:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, data[key])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key in data:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, data[key])

    doc.save(f"docx/{uid}.docx")
