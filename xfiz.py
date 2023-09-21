from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import locale

def fill_contract_xf(uid, a, x):
    locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")
    current_date = datetime.datetime.now()
    formatted_date = current_date.strftime('«%d» %B %Y г.')
    if (x == 1):
        doc = Document("templates/x1fiz.docx")
    else:
        doc = Document("templates/Yearfiz.docx")
    data = {
    f"№______": f"№ {a[0]}",
    "«___»________202_ г.": formatted_date,
    "ТОО «_____________»": f"ТОО «{a[1]}»",
    "БИН:_____________": f"БИН: {a[2]}",
    "директора______________________": f"директора {a[3]}",
    "основании ____________________": f"основании {a[4]}",
    "ФИО________": a[5],
    "ИИН:_____________": f"ИИН: {a[6]}",
    "Удостоверение личности № ________________": f"Удостоверение личности № {a[7]}",
    "от ______________ выдан________________": a[8],
    "адрес: ______________.": f"адрес: {a[9]}",
    "e-mail: _____________.": f"e-mail: {a[10]}",
    "Тел./факс: ____________.": f"Тел./факс: {a[11]}",
    "IBAN: ___________ .": f"IBAN: {a[12]}",
    "в АО __________ .": f"в АО {a[13]}",
    "БИК _____________.": f"БИК {a[14]}",
    "________________________.": a[24],
    "адрес: ______________": f"адрес: {a[15]}",
    "e-mail: _____________": f"e-mail: {a[16]}",
    "Тел./ WhatsApp: ____________": f"Тел./ WhatsApp: {a[17]}",
    "БИН ____________": f"БИН {a[28]}",
    "IBAN: ___________": f"IBAN {a[18]}",
    "в АО __________": f"AO {a[19]}",
    "БИК _____________": f"БИК {a[20]}",
    ", ______________________________.": a[21],
    "в течение ____ (_________________)": f"в течение {a[22]}"
    }
    for paragraph in doc.paragraphs:
        for key in data:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, data[key])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key in data:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, data[key])
    doc.save(f"docx/{uid}.docx")
