from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import locale

def fill_contract_mf(uid, a):
    locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")
    current_date = datetime.datetime.now()
    formatted_date = current_date.strftime('«%d» %B %Y г.')
    text_list = []
    text_list = a[23].split('\n')
    print(text_list)
    doc = Document("templates/mutual_fiz.docx")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith('1.1'):
            p_index = i
            break
    for i, text in enumerate(text_list, start=1):
        new_paragraph = doc.add_paragraph()
        new_paragraph.text = f'1.1.{i} {text}'
        p = new_paragraph._element
        doc.element.body.insert(p_index+i+1, p)
    data = {
    "№______": a[0], #1
    "«___»________202_ г.": formatted_date,
    "ТОО «_____________»": f"ТОО «{a[1]}»", #2
    "БИН:_____________": f"БИН: {a[2]}", #3
    "директора______________________": f"директора {a[3]}", #4
    "основании ____________________": f"основании {a[4]}", #5
    "ФИО________": a[5], #6
    "ИИН:_____________": f"ИИН: {a[6]}",#7
    "Удостоверение личности № ________________": f"Удостоверение личности № {a[7]}",#8
    "от ______________ выдан________________": a[8],#9
    "адрес: ______________.": f"адрес: {a[9]}",#10
    "e-mail: _____________.": f"e-mail: {a[10]}",#11
    "Тел./факс: ____________.": f"Тел./факс: {a[11]}",#12
    "IBAN: ___________ .": f"IBAN: {a[12]}",#13
    "в АО __________ .": f"в АО {a[13]}",#14
    "БИК _____________.": f"БИК {a[14]}",#15
    ".________________________.": a[24],#16
    "адрес: ______________": a[15],#17
    "e-mail: _____________": f"e-mail: {a[16]}",#18
    "Тел./ WhatsApp: ____________": f"Тел./ WhatsApp: {a[17]}",#19
    "БИН ____________": f"БИН {a[28]}",#20
    "IBAN: ___________": f"IBAN {a[18]}",#21
    "в АО __________": f"AO {a[19]}",#22
    "БИК _____________": f"БИК {a[20]}",#23
    "Услуг: ______________________________.": f"Услуг: {a[21]}",#24
    "в течение ____ (_________________)": f"в течение {a[22]}"#25
    }
    for paragraph in doc.paragraphs:
        for key in data:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, data[key])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key in data:
                    if key in cell.text:
                        cell.text = cell.text.replace(key, data[key])
    doc.save(f"docx/{uid}.docx")
